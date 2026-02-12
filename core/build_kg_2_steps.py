import os
import glob
import json
import re
import time
import datetime
import pandas as pd
from tqdm import tqdm
from config_loader import get_config
from core.llm.utils import parse_json_response
from kg.prompts import (
    DIET_KG_EXTRACT_COT_PROMPT_v1,
    DIET_KG_RESOLUTION_PROMPT_v1,
    EXER_KG_EXTRACT_COT_PROMPT_v1,
    EXER_KG_RESOLUTION_PROMPT_v1
)
# Optional import for local model support
try:
    from core.llm import should_use_local, get_unified_llm
    HAS_UNIFIED_LLM = True
except ImportError:
    HAS_UNIFIED_LLM = False
# Libraries for handling PDF and Word
import pymupdf4llm
from docx import Document

config = get_config()
API_MODEL = config.get("api_model", {})
DEEPSEEK_API_KEY = API_MODEL.get("api_key", "")
DEEPSEEK_BASE_URL = API_MODEL.get("base_url", "")
MODEL_NAME = API_MODEL.get("model", "deepseek-chat")
# Always use API model (no local model fallback)
USE_LOCAL = False
KG_EXTRACT_LOG_PATH = "kg/llm_log.log"
print(f"[INFO] KG Builder LLM mode: api")
print(f"[INFO] API Model: {MODEL_NAME} @ {DEEPSEEK_BASE_URL}")

# Knowledge Graph Type Configuration
KG_CONFIGS = {
    "diet": {
        "input_dir": "data/diet",
        "name": "Diet",
        "use_two_step": True,
        "cot_prompt": DIET_KG_EXTRACT_COT_PROMPT_v1,
        "resolution_prompt": DIET_KG_RESOLUTION_PROMPT_v1
    },
    "exercise": {
        "input_dir": "data/exer",
        "name": "Exercise",
        "use_two_step": True,
        "cot_prompt": EXER_KG_EXTRACT_COT_PROMPT_v1,
        "resolution_prompt": EXER_KG_RESOLUTION_PROMPT_v1
    }
}


# Base directory for saving results (all history records will be stored in this folder)
OUTPUT_BASE_DIR = "output_history"
# Text splitting settings
CHUNK_SIZE = 1000
OVERLAP = 200



def read_excel(file_path):
    """
    [New] Read Excel and convert each row into natural language sentences
    """
    print(f"Parsing Excel: {os.path.basename(file_path)}")
    text_content = []
    try:
        # Read all worksheets (sheet_name=None returns a dictionary)
        dfs = pd.read_excel(file_path, sheet_name=None, engine='openpyxl')
        for sheet_name, df in dfs.items():
            if df.empty: continue
            # 1. Clean headers (convert to strings, remove spaces)
            headers = [str(col).strip().replace("\n", "") for col in df.columns]
            # 2. Iterate through each row
            # fillna('') to prevent errors from empty values
            for _, row in df.fillna('').iterrows():
                row_parts = []
                for header, cell_value in zip(headers, row):
                    # If the cell is not empty, concatenate "header is value"
                    val_str = str(cell_value).strip().replace("\n", " ")
                    if val_str and val_str.lower() != 'nan':
                        row_parts.append(f"{header} is {val_str}")
                # 3. Combine into a sentence
                if row_parts:
                    # Example: "In the data table Sheet1, drug is metformin, dosage is 500mg."
                    sentence = f"In the data table {sheet_name}, " + ", ".join(row_parts) + "."
                    text_content.append(sentence)
        return "\n".join(text_content)
    except Exception as e:
        print(f"Excel read failed {file_path}: {e}")
        return ""


def read_docx(file_path):
    """ Extract Word, including table to natural language logic """
    try:
        doc = Document(file_path)
        text_content = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        if doc.tables:
            for table in doc.tables:
                headers = [cell.text.strip().replace("\n", "") for cell in table.rows[0].cells]
                for row in table.rows[1:]:
                    row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    row_parts = []
                    for i in range(len(row_cells)):
                        if i < len(headers) and row_cells[i]:
                            row_parts.append(f"{headers[i]} is {row_cells[i]}")
                    if row_parts:
                        text_content.append(", ".join(row_parts) + ".")
        return "\n".join(text_content)
    except Exception as e:
        print(f"Word read failed {file_path}: {e}")
        return ""


def read_pdf(file_path):
    """ Extract PDF (pymupdf4llm) """
    try:
        return pymupdf4llm.to_markdown(file_path)
    except Exception as e:
        print(f"PDF read failed {file_path}: {e}")
        return ""


def read_txt(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except:
        return ""


def clean_text(text):
    text = re.sub(r'\[\d+(?:,\s*\d+)*\]', '', text)
    text = re.sub(r'\n\s*\d+\s*\n', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def split_text_by_headers(text, chunk_size=CHUNK_SIZE):
    """Split text by Markdown headers (##) to keep sections together."""
    if not text: return []
    sections = re.split(r'(^##\s+.*)', text, flags=re.MULTILINE)
    chunks = []
    current_chunk = ""
    for part in sections:
        if not part: continue
        if len(current_chunk) + len(part) > chunk_size:
            if current_chunk.strip():
                chunks.append(current_chunk.strip())
            current_chunk = part
        else:
            current_chunk += part
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    return chunks


def split_text(text, chunk_size=CHUNK_SIZE, overlap=OVERLAP):
    """Legacy fallback: simple chunking by character limit."""
    if not text: return []
    chunks = []
    start = 0
    length = len(text)
    while start < length:
        end = min(start + chunk_size, length)
        if end < length:
            next_newline = text.find('\n', end, end + 100)
            if next_newline != -1:
                end = next_newline
        chunk = text[start:end]
        if len(chunk.strip()) > 20:
            chunks.append(chunk)
        start += (chunk_size - overlap)
    return chunks


def _call_llm(prompt, temperature=0.1):
    """Helper function to call LLM with a prompt."""
    # Log the prompt
    _log_llm_interaction("PROMPT", prompt)

    messages = [
        {"role": "system", "content": "You are a helpful medical assistant. Always output valid JSON."},
        {"role": "user", "content": prompt}
    ]
    from openai import OpenAI
    client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url=DEEPSEEK_BASE_URL)
    response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=messages,
        temperature=temperature,
        stream=False,
        response_format={'type': 'json_object'}
    )
    # lower to remove duplicates in Letter Cases
    content = response.choices[0].message.content.strip().lower()

    # Log the response
    _log_llm_interaction("RESPONSE", content)

    return content


def _log_llm_interaction(step, content):
    """Log LLM prompts and responses to the log file."""
    try:
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S.%f")
        with open(KG_EXTRACT_LOG_PATH, 'a', encoding='utf-8') as f:
            f.write(f"\n{'='*60}\n")
            f.write(f"[{timestamp}] {step}\n")
            f.write(f"{'='*60}\n")
            f.write(f"{content}\n")
            f.write(f"\n{'='*60}\n")
    except Exception as e:
        print(f"[WARN] Failed to log LLM interaction: {e}")


def _apply_canonical_mapping(quads, entity_mapping):
    """
    Apply canonical entity mapping to quads.

    Args:
        quads: List of quad dictionaries with 'head', 'relation', 'tail', 'context'
        entity_mapping: Dict mapping original entities to canonical forms

    Returns:
        List of quads with entities replaced by canonical forms
    """
    if not entity_mapping:
        return quads

    cleaned_quads = []
    for quad in quads:
        new_head = entity_mapping.get(quad.get('head', ''), quad.get('head', ''))
        new_tail = entity_mapping.get(quad.get('tail', ''), quad.get('tail', ''))
        cleaned_quads.append({
            'head': new_head,
            'relation': quad.get('relation', ''),
            'tail': new_tail,
            'context': quad.get('context', 'General')
        })
    return cleaned_quads


def _resolve_entities(extracted_entities, resolution_prompt):
    """
    Call LLM to resolve duplicate entities and get canonical forms.

    Args:
        extracted_entities: List of entity strings
        resolution_prompt: The resolution prompt template

    Returns:
        Dict mapping original entities to canonical forms, or empty dict on failure
    """
    if not extracted_entities or not resolution_prompt:
        return {}

    # Format entities as a comma-separated list
    entities_list = ", ".join([f'"{e}"' for e in extracted_entities])
    # prompt = f"{resolution_prompt}\n\n## Extracted Entities\n[{entities_list}]"
    prompt = resolution_prompt(ENTITIES=entities_list)

    try:
        content = _call_llm(prompt, temperature=0.2)
        data = parse_json_response(content)

        if isinstance(data, dict) and "resolutions" in data:
            resolutions = data["resolutions"]
            if not resolutions:
                return {}

            # Build mapping from duplicates to canonical form
            entity_mapping = {}
            for resolution in resolutions:
                canonical = resolution.get("canonical_form", "")
                duplicates = resolution.get("duplicate_group", [])
                for dup in duplicates:
                    if dup != canonical:
                        entity_mapping[dup] = canonical
            return entity_mapping
        return {}
    except Exception as e:
        print(f"[WARN] Entity resolution failed: {e}")
        return {}


def extract_quads_with_llm(text_chunk, cot_prompt=None, resolution_prompt=None, use_two_step=True):
    """
    Extract knowledge quads from text chunk using two-step workflow.

    Two-step workflow (when use_two_step=True):
    1. Extract entities + quads using Chain of Thought prompt
    2. Resolve entity duplicates using resolution prompt
    3. Apply canonical forms to quads

    Args:
        text_chunk: Text content to process
        cot_prompt: Chain of Thought extraction prompt (for step 1)
        resolution_prompt: Entity resolution prompt (for step 2)
        use_two_step: Whether to use two-step workflow

    Returns:
        List of quads with standardized entities
    """
    if len(text_chunk.strip()) < 10:
        return []

    # Fallback to one-step if two-step not enabled or prompts not available
    if not use_two_step or not cot_prompt:
        return []

    # ========== STEP 1: Extract entities and quads using CoT ==========
    cot_prompt_text = cot_prompt(TEXT=text_chunk)
    try:
        content = _call_llm(cot_prompt_text, temperature=0.1)
        data = parse_json_response(content)

        if not isinstance(data, dict):
            print(f"[WARN] CoT extraction returned non-dict")
            return []

        # Extract entities and quads from CoT response
        extracted_entities = data.get("extracted_entities", [])
        quads = data.get("quads", [])

        if not quads:
            print(f"[WARN] No quads extracted from chunk")
            return []

        print(f"[DEBUG] Step 1: Extracted {len(quads)} quads and {len(extracted_entities)} entities")

    except json.JSONDecodeError as e:
        print(f"[WARN] CoT JSON parsing failed: {e}")
        return []
    except Exception as e:
        print(f"[WARN] CoT extraction failed: {e}")
        return []

    # ========== STEP 2: Resolve entity duplicates ==========
    if extracted_entities and resolution_prompt:
        entity_mapping = _resolve_entities(extracted_entities, resolution_prompt)

        if entity_mapping:
            print(f"[DEBUG] Step 2: Resolved {len(entity_mapping)} entity mappings")
            quads = _apply_canonical_mapping(quads, entity_mapping)

    return quads


def build_knowledge_graph(kg_type: str, config: dict) -> dict:
    """
    Build knowledge graph for a specific type (diet or exercise).
    Args:
        kg_type: Type of knowledge graph ('diet' or 'exercise')
        config: Configuration dict with input_dir, name, use_two_step
    Returns:
        Dict with stats about the build
    """
    kg_name = config["name"]
    input_dir = config["input_dir"]
    use_two_step = config.get("use_two_step", False)
    cot_prompt = config.get("cot_prompt")
    resolution_prompt = config.get("resolution_prompt")
    # Check input folder
    if not os.path.exists(input_dir):
        os.makedirs(input_dir)
        print(f"[{kg_name} KG] Please create {input_dir} and put files in it")
        return {"status": "skipped", "reason": "input_dir_not_found", "quads": 0}
    # Generate output folder for this run
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    current_output_dir = os.path.join(OUTPUT_BASE_DIR, f"{kg_type.capitalize()}_{timestamp}")
    # Create folder
    os.makedirs(current_output_dir, exist_ok=True)
    print(f"[{kg_name} KG] This result will be saved in: {current_output_dir}")
    # Scan files
    files = glob.glob(os.path.join(input_dir, "*.pdf")) + \
        glob.glob(os.path.join(input_dir, "*.docx")) + \
        glob.glob(os.path.join(input_dir, "*.txt")) + \
        glob.glob(os.path.join(input_dir, "*.xlsx"))
    if not files:
        print(f"[{kg_name} KG] '{input_dir}' folder is empty, no files found.")
        return {"status": "skipped", "reason": "no_files", "quads": 0}
    # Load checkpoint CSV to skip already processed files
    checkpoint_csv_path = os.path.join(OUTPUT_BASE_DIR, f"{kg_type}_processed_files.csv")
    processed_files_checkpoint = set()
    if os.path.exists(checkpoint_csv_path):
        try:
            checkpoint_df = pd.read_csv(checkpoint_csv_path)
            processed_files_checkpoint = set(checkpoint_df["file_path"].tolist())
            print(f"[{kg_name} KG] Loaded checkpoint: {len(processed_files_checkpoint)} files already processed")
        except Exception as e:
            print(f"[{kg_name} KG] Failed to load checkpoint: {e}")
    # Filter out already processed files
    new_files = [f for f in files if f not in processed_files_checkpoint]
    skipped_count = len(files) - len(new_files)
    if skipped_count > 0:
        print(f"[{kg_name} KG] Skipping {skipped_count} already processed files")
    files = new_files
    print(f"[{kg_name} KG] Found {len(files)} new files to process...")
    if not files:
        print(f"[{kg_name} KG] All files have been processed. Nothing to do.")
        return {"status": "skipped", "reason": "all_files_processed", "quads": 0}
    all_quads = []
    seen_hashes = set()
    processed_files_log = []
    start_time = time.time()
    # Loop processing
    files_processed_this_run = []
    for file_path in tqdm(files, desc=f"{kg_name} KG Progress"):
        file_name = os.path.basename(file_path)
        processed_files_log.append(file_name)
        ext = file_path.lower()
        if ext.endswith(".pdf"): content = read_pdf(file_path)
        elif ext.endswith(".docx"): content = read_docx(file_path)
        elif ext.endswith(".xlsx"): content = read_excel(file_path)
        else: content = read_txt(file_path)
        if not content:
            # Mark file as processed even if no content
            files_processed_this_run.append({
                "file_path": file_path,
                "file_name": file_name,
                "processed_at": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "status": "no_content"
            })
            continue
        # Clean text and split by headers
        cleaned_content = clean_text(content)
        chunks = split_text_by_headers(cleaned_content)
        for chunk in tqdm(chunks, desc=f"Parsing {file_name[:10]}", leave=False):
            quads = extract_quads_with_llm(
                chunk,
                cot_prompt=cot_prompt,
                resolution_prompt=resolution_prompt,
                use_two_step=use_two_step
            )
            for t in quads:
                if "head" in t and "relation" in t and "tail" in t:
                    # Include context in hash for deduplication
                    context = t.get('context', 'General')
                    t_hash = f"{t['head']}_{t['relation']}_{t['tail']}_{context}"
                    if t_hash not in seen_hashes:
                        seen_hashes.add(t_hash)
                        t["source"] = file_name
                        all_quads.append(t)
                time.sleep(0.1)
        # Mark file as successfully processed and update checkpoint immediately
        files_processed_this_run.append({
            "file_path": file_path,
            "file_name": file_name,
            "processed_at": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "status": "success"
        })
        # Append to checkpoint CSV
        try:
            checkpoint_df_update = pd.DataFrame([files_processed_this_run[-1]])
            if os.path.exists(checkpoint_csv_path):
                checkpoint_df_update.to_csv(checkpoint_csv_path, mode='a', header=False, index=False)
            else:
                checkpoint_df_update.to_csv(checkpoint_csv_path, index=False)
        except Exception as e:
            print(f"Failed to update checkpoint for {file_name}: {e}")
    # Save results
    duration = time.time() - start_time
    output_json_path = os.path.join(current_output_dir, f"{kg_type}_quads.json")
    output_csv_path = os.path.join(current_output_dir, f"{kg_type}_quads.csv")
    log_path = os.path.join(current_output_dir, "process_log.txt")
    print("-" * 40)
    print(f"[{kg_name} KG] Extraction complete! Time: {duration:.2f} seconds")
    print(f"Obtained {len(all_quads)} unique quads.")
    # Save JSON
    with open(output_json_path, 'w', encoding='utf-8') as f:
        json.dump(all_quads, f, indent=4, ensure_ascii=False)
    # Save CSV
    df = pd.DataFrame(all_quads)
    if not df.empty:
        cols = ["head", "relation", "tail", "context", "source"]
        existing = [c for c in cols if c in df.columns]
        df = df[existing]
        df.to_csv(output_csv_path, index=False, encoding='utf_8_sig')
    # Save log
    with open(log_path, 'w', encoding='utf-8') as f:
        f.write(f"Run time: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"Knowledge graph type: {kg_name}\n")
        f.write(f"Time consumed: {duration:.2f} seconds\n")
        f.write(f"Number of extracted quads: {len(all_quads)}\n")
        f.write(f"Input directory: {input_dir}\n")
        f.write("\nList of processed files:\n")
        for fname in processed_files_log:
            f.write(f"- {fname}\n")
    print(f"[{kg_name} KG] Results saved to: {current_output_dir}")
    return {
        "status": "success",
        "kg_type": kg_type,
        "kg_name": kg_name,
        "quads": len(all_quads),
        "duration": duration,
        "output_dir": current_output_dir
    }


def main():
    """Build both diet and exercise knowledge graphs by default."""
    # Check command line arguments
    kg_types_to_build = None
    for arg in sys.argv[1:]:
        if arg.startswith("--kg="):
            kg_types_to_build = [arg.replace("--kg=", "").lower()]
            break
        elif arg in ["-h", "--help"]:
            print("""
Usage: python -m core.build_kg [options]
Options:
  --kg=diet Build only diet knowledge graph
  --kg=exercise Build only exercise knowledge graph
  --kg=all Build diet and exercise knowledge graphs (default behavior)
  -h, --help Display this help information
Default behavior:
  If no options are specified, build both diet and exercise knowledge graphs.
            """)
            return
    # If not specified, build both KGs by default
    if kg_types_to_build is None or kg_types_to_build[0] == "all":
        kg_types_to_build = ["diet", "exercise", "others"]
    print("=" * 50)
    print(f"Starting to build knowledge graph...")
    print(f"Types: {', '.join(kg_types_to_build)}")
    print("=" * 50)
    total_stats = {
        "total_quads": 0,
        "total_duration": 0,
        "results": []
    }
    for kg_type in kg_types_to_build:
        if kg_type in KG_CONFIGS:
            print()
            stats = build_knowledge_graph(kg_type, KG_CONFIGS[kg_type])
            total_stats["results"].append(stats)
            total_stats["total_quads"] += stats.get("quads", 0)
            total_stats["total_duration"] += stats.get("duration", 0)
        else:
            print(f" Unknown knowledge graph type: {kg_type}")
    # Summary
    print()
    print("=" * 50)
    print("Knowledge graph build summary")
    print("=" * 50)
    print(f"Total quads count: {total_stats['total_quads']}")
    print(f"Total time: {total_stats['total_duration']:.2f} seconds")
    print()
    # Display status for each KG
    for stats in total_stats["results"]:
        status = "[pass]" if stats.get("status") == "success" else "[error]"
        print(f" {status} {stats.get('kg_name', 'Unknown')} KG: {stats.get('quads', 0)} quads")
    print()
    print("File locations:")
    for stats in total_stats["results"]:
        if stats.get("output_dir"):
            print(f" - {stats.get('kg_name', '')}: {stats['output_dir']}")

import sys
if __name__ == "__main__":
    main()
